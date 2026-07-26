"""
Tender Trimesters — Marketing Assets (DOCX + XLSX)
Generates:
1. social-content-kit.docx
2. launch-announcement.docx
3. influencer-outreach.docx
4. content-calendar.xlsx
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

OUTPUT_DIR = "/home/z/my-project/download"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Brand colors (RGB)
MOSS_DEEP = RGBColor(0x3A, 0x42, 0x33)
MOSS = RGBColor(0x6B, 0x7A, 0x5A)
TERRACOTTA = RGBColor(0xC9, 0x7B, 0x5C)
BLUSH = RGBColor(0xFA, 0xDA, 0xDD)
ROSE_GOLD = RGBColor(0xB7, 0x6E, 0x79)
CREAM = RGBColor(0xF5, 0xEF, 0xE0)
INK = RGBColor(0x2D, 0x2A, 0x24)

# Excel hex fills
HEX_MOSS_DEEP = "3A4233"
HEX_MOSS = "6B7A5A"
HEX_CREAM = "F5EFE0"
HEX_BLUSH = "FADADD"
HEX_SAGE = "CFE3DC"
HEX_BUTTER = "FFF6E5"

def shade_cell(cell, hex_color):
    """Shade a table cell with the given hex color."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)

def add_horizontal_line(doc, color=ROSE_GOLD):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B76E79")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

def set_paragraph_font(paragraph, font_name="Montserrat", size=11, color=INK, bold=False, italic=False):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.font.bold = bold
        run.font.italic = italic

def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Georgia"  # fallback for serif
    run.font.size = Pt(28 if level == 1 else 22 if level == 2 else 16)
    run.font.color.rgb = MOSS_DEEP
    run.font.bold = True
    return p

def add_para(doc, text, size=11, color=INK, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, font="Calibri"):
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    return p

# ──────────────────────────────────────────────────────────────────────────
# 1. SOCIAL CONTENT KIT
# ──────────────────────────────────────────────────────────────────────────

def generate_social_content_kit():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TENDER TRIMESTERS")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = ROSE_GOLD
    run.font.bold = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run("Social Content Kit")
    run.font.name = "Georgia"
    run.font.size = Pt(36)
    run.font.color.rgb = MOSS_DEEP
    run.font.bold = True
    add_para(doc, "30 ready-to-post captions  ·  5 hashtag sets  ·  Weekly posting rhythm",
             size=11, color=MOSS, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_horizontal_line(doc)
    add_para(doc, "by Mommies Matter  ·  Made with love, mama",
             size=9, color=MOSS, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # How to use
    add_heading_styled(doc, "How to use this kit", level=2)
    add_para(doc, "This kit gives you 30 days of ready-to-post content for Instagram and TikTok. Each caption is written in the Tender Trimesters voice — warm, nurturing, like a best friend who happens to know a lot about pregnancy. Mix and match. Edit to make them yours. The goal is consistency, not perfection.",
             size=11)
    add_para(doc, "Posting rhythm: 4–5 posts per week is sustainable for most solo founders. Start there. Each caption below includes a hook (the first line that stops the scroll), the body, and a call-to-action. Paste the hook as your on-screen text for TikTok/Reels. Use the body as your caption.",
             size=11)
    add_para(doc, "Hashtag sets are at the end. Rotate them — don't copy-paste the same 30 hashtags every post. The algorithm prefers variety.",
             size=11)

    doc.add_page_break()

    # Captions
    add_heading_styled(doc, "30 Captions", level=2)

    captions = [
        # Educational
        ("educational", "Week 12 — what's actually happening",
         "Week 12. Your baby has reflexes now.",
         "If you poke your belly, baby might squirm in response. All their systems are forming. The NT scan (nuchal translucency) happens around now — it checks for chromosomal conditions and confirms your due date.\n\nThis is also the week many mamas feel safe enough to announce. There's no 'right' time. Do what feels true to you.\n\nWhat week are you in? Drop it below 👇",
         "Save this for when you hit week 12."),

        ("educational", "The 5-1-1 rule for contractions",
         "When do you actually go to the hospital?",
         "It's the 5-1-1 rule:\n• Contractions 5 minutes apart\n• Each lasting 1 minute\n• For at least 1 hour\n\nCall your OB sooner if your water breaks, you have heavy bleeding, or baby's movement decreases.\n\nSave this. Send to your birth partner. You'll thank us later.",
         "Save this for labor day."),

        ("educational", "First trimester foods to avoid",
         "Foods to skip in your first trimester:",
         "Skip: raw sushi, deli meats (unless heated), soft unpasteurized cheeses, high-mercury fish (swordfish, king mackerel), raw sprouts, undercooked eggs.\n\nLimit: caffeine to under 200mg/day (one 12oz coffee).\n\nEat: protein, leafy greens, omega-3s, anything that sounds good. Calories matter more than perfection right now.\n\nWhat's your weirdest craving?",
         "Save + share with a mama who needs this."),

        ("emotional", "You're not 'too much'",
         "If you've cried at a commercial this week —",
         "That's your hormones. Not you being 'too much.'\n\nFirst trimester mood swings are real. Your body is producing estrogen and progesterone at levels you've never experienced. You are allowed to feel everything. All at once. Without apology.\n\nYou're not alone, mama.",
         "Send this to a mama who needs to hear it today."),

        ("emotional", "Letter to week 16 you",
         "To the mama at week 8 wondering if she'll ever feel 'glowing' —",
         "You will.\n\nRight now you're tired. Nauseous. Maybe anxious. Maybe overjoyed. Often all of those before lunch. That's first trimester.\n\nSecond trimester is coming. Energy returns. The glow appears. You'll feel baby move for the first time. You'll look pregnant and not just 'did you eat too much pizza.'\n\nHang in there. We see you.",
         "Save for your 3am panic."),

        ("product", "Meet Tempie — your 24/7 AI companion",
         "It's 3am. You're anxious. Your OB is asleep.",
         "Tempie isn't.\n\nShe's the AI companion inside Tender Trimesters — trained to listen without judgment, offer practical tips, share affirmations, and remind you that you're not alone. She's not a doctor. She's your person at 3am.\n\nFree tier: 5 messages/day. Premium: unlimited.\n\nTry her free at the link in bio.",
         "Link in bio to try Tempie free."),

        ("product", "Free vs Premium — what's the difference?",
         "Free, or fully held?",
         "FREE (forever):\n• 40-week milestone calendar\n• Daily affirmations\n• Mood tracking\n• Private journal\n• Tempie — 5 msgs/day\n\nPREMIUM ($9.99 one-time):\n• Unlimited Tempie\n• Bump photo gallery\n• Partner access\n• Audio meditations\n• Kick counter\n• The full Mommies Matter bundle (ebook + affirmation deck + letters to baby templates)\n\nStart free. Upgrade when you're ready. Link in bio.",
         "Comment 'PREMIUM' for the link."),

        ("bts", "Why I built this",
         "I built Tender Trimesters because —",
         "When I was pregnant with my son Tatum, I needed someone to tell me what to actually do. Not the textbook version. The real version.\n\nThe one that says: eat the saltines, take the nap, call your OB when you're worried, and trust that you are not alone in this.\n\nSo I built it. For you. For the mama I was.\n\nWith love, Helena-Ann 💛",
         "Follow for the journey."),

        ("bts", "Behind the affirmation cards",
         "Each affirmation in the app is written by hand.",
         "Not AI. Not a stock list. Real words I needed to hear when I was pregnant with Tatum.\n\n40 of them — one for each week. From 'My body knows exactly what to do' (week 1) to 'Today could be the day. I am ready' (week 40).\n\nDownload the app for the full deck. Or comment 'CARDS' and we'll send you a printable sample.",
         "Comment CARDS for a free sample."),

        ("affirmation", "Today's affirmation",
         "Your affirmation for today:",
         "\"My changing body is beautiful.\"\n\nSay it out loud. Let it sink in. Repeat as needed.\n\nThis is week 16's affirmation in the Tender Trimesters app. 40 weeks of them. One for each week of your pregnancy.\n\nWhat's your affirmation today?",
         "Save + repeat daily."),

        ("affirmation", "Monday affirmation",
         "Mondays are hard. This helps.",
         "\"Rest is productive. My body is doing important work.\"\n\nPinned to your mirror. Set as your phone wallpaper. Tattooed on your forearm (kidding, mostly).\n\nHave a gentle Monday, mama.",
         "Save for the Mondays that need it."),

        ("ebook", "From the book — should I get an epidural?",
         "Excerpt from 'Mommies Matter' by Helena-Ann Baker —",
         "\"Birth is unpredictable. But that doesn't mean you're powerless.\n\nI was 18 hours in. My body was tired. I had been moaning through contractions and passing out in between them without even realizing it. I could barely keep my eyes open. At 7 centimeters, I cried to my husband: 'I can't do it, baby. I need the epidural.'\n\nSuddenly the nurses were in motion — like it was a national emergency...\"\n\nRead the full chapter in Mommies Matter — live on Amazon Kindle.",
         "Link in bio to read the first chapter free."),

        ("ebook", "Helena-Ann's birth story — preeclampsia + Tatum",
         "I almost didn't write this chapter.",
         "It's the hardest one. My birth story with Tatum. Preeclampsia. Pain. The moments I thought I might not make it. The moment I first held him.\n\nI wrote it because someone needs to read it. Maybe you. Maybe your best friend. Maybe the mama at 2am googling 'preeclampsia birth story' the way I did.\n\nYou're not alone in the hard parts. That's why I wrote the book.\n\nMommies Matter — live on Amazon Kindle.",
         "Link in bio."),

        ("educational", "What to pack in your hospital bag",
         "Hospital bag essentials (that nobody tells you):",
         "• Long phone charger (10ft) — outlets are never close\n• Flip flops for the shower\n• Your own pillow (hospital ones are terrible)\n• Lip balm (hospitals are dry)\n• Snacks for your partner (vending machines lie)\n• A going-home outfit that's basically pajamas\n• Baby's outfit (one in newborn size, one in 0-3mo — you don't know)\n• Insurance card + ID\n• Hair ties\n\nWhat did we miss? Drop your hospital bag must-haves below.",
         "Save for week 36."),

        ("educational", "What's a kick count?",
         "Kick counts — what they are and when to start:",
         "Start around week 28. Once a day, ideally after a meal when baby's active.\n\nLie on your left side. Count kicks until you reach 10. Should take less than 2 hours.\n\nIf you don't feel 10 kicks in 2 hours, call your OB. Not a panic — a phone call. They'd rather hear from you.\n\nSave this for week 28.",
         "Save + share with a third-trimester mama."),

        ("emotional", "The two-week wait",
         "The two-week wait is its own kind of marathon.",
         "You're not technically pregnant yet. Or maybe you are — you won't know for another week.\n\nEvery twinge is a sign. Every non-twinge is a sign. Google is your enemy. Hope is your frenemy.\n\nIf you're in the TWW right now — we see you. Be gentle with yourself today.",
         "Send to a friend in her TWW."),

        ("emotional", "Pregnancy after loss",
         "Pregnancy after loss is its own kind of brave.",
         "Every appointment feels like a verdict. Every week that passes is a small victory. Every ultrasound is hope mixed with terror.\n\nIf this is you — we hold space for you here. Your joy is allowed to be complicated. Your grief is allowed to coexist with your hope.\n\nYou are not alone. 💛",
         "Share with someone who needs this."),

        ("product", "The bundle is live",
         "Everything mama needs, in one place.",
         "Premium Bundle ($9.99 one-time, was $29.99):\n• Mommies Matter ebook — 17 chapters\n• 40 printable affirmation cards\n• First-Trimester Survival Kit\n• Letters to Baby journal templates\n• Unlimited Tempie AI chat\n• Bump photo gallery\n• Partner access\n• Audio meditations\n• Kick counter\n\nOr $4.99/mo if you prefer to spread it out.\n\nLink in bio.",
         "Link in bio."),

        ("bts", "Why we don't use baby pink",
         "Hot take: pregnancy apps don't have to be baby pink.",
         "When we designed Tender Trimesters, we chose earthy moss, dusty terracotta, hearth cream — with blush and rose gold accents.\n\nWhy? Because pregnancy is sacred, not saccharine. Because mamas deserve an aesthetic that feels like a sanctuary, not a candy store.\n\nThe blush is still there — we love soft pink. But it's a supporting character, not the whole movie.",
         "Do you prefer earthy or all-blush? Vote below."),

        ("affirmation", "Affirmation for anxiety",
         "For the mama spiraling at 3am:",
         "\"My worries are valid, and they are not in charge.\"\n\nSay it. Breathe it. Write it on your hand if you need to.\n\nIf anxious thoughts are keeping you up, write them down. Getting them out of your head and onto paper helps. So does Tempie — she's awake when you are.",
         "Save + send to a 3am friend."),

        ("ebook", "What I wish I'd known",
         "What I wish I'd known before becoming a mom:",
         "1. The 'fourth trimester' is real. Newborn days are survival mode.\n2. Asking for help is not weakness. It's strategy.\n3. Your relationship with your partner will change. Tend to it.\n4. Sleep when the baby sleeps is annoying advice but it's true.\n5. You will cry over spilled milk. Both literally and metaphorically.\n6. You will also feel a love you didn't know was possible.\n\nRead the full chapter in Mommies Matter — Amazon Kindle.",
         "Link in bio."),

        ("educational", "Glucose test — what to expect",
         "Glucose test week (24-28). Here's the real talk:",
         "You drink a sweet orange (or fruit punch, yum) drink. Wait an hour. Get your blood drawn. That's it.\n\nTips:\n• Schedule it for first thing in the morning\n• Eat a protein breakfast (eggs, not cereal)\n• Bring a book or podcast — the hour feels long\n• If you fail the 1-hour, you do the 3-hour. Not the end of the world.\n\nGestational diabetes is manageable. You've got this.",
         "Save for week 24-28."),

        ("emotional", "Body image in pregnancy",
         "Your body is doing extraordinary work.",
         "Some days you'll feel like Superwoman. Other days you'll feel like a stranger in your own skin.\n\nBoth are normal. The linea nigra, the stretch marks, the swelling, the pregnancy mask — all temporary. All badges of the work your body is doing.\n\nBe gentle with her. She's making a human.",
         "Save for the hard body days."),

        ("product", "Partner access — bring them along",
         "Your partner wants to be involved. They don't always know how.",
         "Tender Trimesters Premium includes partner access — a read-only link to your pregnancy journey. They see your current week, this week's content, upcoming appointments, and a 'how to support her' tip.\n\nThey can't see your journal. They can't chat with Tempie. Just the journey.\n\nGenerate your link in Profile > Partner access.",
         "Tag your partner below."),

        ("bts", "Helena-Ann's why",
         "Why I started Mommies Matter —",
         "I lost my Granny in 2019. She came to me in a dream and told me I was pregnant before I even knew.\n\nShe had asked me for years: 'Helena-Ann, when you gonna write your book?' She saw the writer in me before I did.\n\nI believe she met Tatum before I did, and sent him to me because she knew I would need his love after she left.\n\nThis is for her. Always. And for every mama who feels unsure, unseen, or undone. You matter.",
         "Follow for the journey."),

        ("affirmation", "Friday affirmation",
         "Friday affirmation:",
         "\"I trust the timing of my life.\"\n\nWhether you're trying to conceive, in your first trimester, or 39 weeks and over it — the timing of your life is unfolding exactly as it should.\n\nHave a gentle weekend, mama.",
         "Save for Fridays."),

        ("educational", "The anatomy scan — week 20",
         "Week 20 anatomy scan — what they're actually checking:",
         "It's the big ultrasound. They check:\n• Baby's organs (heart, brain, spine, kidneys, stomach)\n• Growth measurements\n• Placenta position\n• Amniotic fluid levels\n• Sometimes, the sex (if you want to know)\n\nTakes 30-45 minutes. Bring your partner. Ask for printed photos.\n\nYou're halfway there, mama.",
         "Save for week 20."),

        ("emotional", "To the mama who feels behind",
         "If you feel like you're 'behind' in life —",
         "Married at 25. House at 28. Baby at 30. The timeline was never real.\n\nYou're not behind. You're on your own timeline. The love you'll give this baby doesn't care about your age, your zip code, or your savings balance.\n\nYou are exactly where you're supposed to be.",
         "Send to a mama who needs this."),

        ("product", "Free forever",
         "Tender Trimesters is free forever.",
         "Not a 14-day trial. Not 'free with credit card required.' Free, forever.\n\n40-week milestone calendar. Daily affirmations. Mood tracking. Private journal. Tempie AI chat — 5 messages a day.\n\nWe made it free because every mama deserves support, regardless of budget. Premium ($9.99 one-time) is there when you want more — but you'll never be locked out of the essentials.\n\nLink in bio.",
         "Link in bio."),

        ("bts", "Tempie's training",
         "How we trained Tempie to be a 'she' not an 'it':",
         "Tempie's personality is built on three principles:\n\n1. Warm, not clinical. Like a best friend who happens to know a lot about pregnancy.\n2. Listen without judgment. Don't fix everything. Sometimes mamas just need to be heard.\n3. Always defer to her OB for medical stuff. Tempie's not a doctor. She's a friend who knows when to say 'call your provider.'\n\nThe result is an AI companion that feels like talking to a real person at 3am. That was the goal.",
         "Try Tempie free — link in bio."),
    ]

    for i, (category, hook_title, hook, body, cta) in enumerate(captions, 1):
        # Caption number + category badge
        p = doc.add_paragraph()
        run = p.add_run(f"#{i:02d}  ·  {category.upper()}")
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = ROSE_GOLD
        run.font.bold = True

        # Hook title
        p = doc.add_paragraph()
        run = p.add_run(hook_title)
        run.font.name = "Georgia"
        run.font.size = Pt(16)
        run.font.color.rgb = MOSS_DEEP
        run.font.bold = True

        # Hook (first line)
        p = doc.add_paragraph()
        run = p.add_run(hook)
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        run.font.color.rgb = INK
        run.font.bold = True

        # Body
        p = doc.add_paragraph()
        run = p.add_run(body)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = INK

        # CTA
        p = doc.add_paragraph()
        run = p.add_run(f"CTA: {cta}")
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = TERRACOTTA
        run.font.italic = True

        # Separator
        add_horizontal_line(doc)

    doc.add_page_break()

    # Hashtag sets
    add_heading_styled(doc, "Hashtag Sets", level=2)
    add_para(doc, "Rotate these. Don't copy-paste the same set every post. The algorithm prefers variety. Mix 8–12 hashtags per post, with a combination of broad (#NewMom) and specific (#TenderTrimesters).",
             size=11, italic=True)

    hashtag_sets = [
        ("Set 1: Brand & App", ["#TenderTrimesters", "#MommiesMatter", "#PregnancyApp", "#PregnancyJourney", "#NewMomLife", "#ExpectingMom", "#PregnancyCalendar", "#MamaToBe", "#HelenaAnnBaker", "#TempieAI"]),
        ("Set 2: First Trimester", ["#FirstTrimester", "#EarlyPregnancy", "#PregnancySymptoms", "#MorningSickness", "#PregnancyAnnouncement", "#12WeeksPregnant", "#PregnancyTips", "#FirstTimeMom", "#TryingToConceive", "#TWW"]),
        ("Set 3: Second Trimester", ["#SecondTrimester", "#BabyBump", "#PregnancyGlow", "#20WeeksPregnant", "#AnatomyScan", "#PregnantLife", "#BumpPic", "#PregnancyAffirmations", "#MamaLife", "#TenderTrimesters"]),
        ("Set 4: Third Trimester", ["#ThirdTrimester", "#CountingDown", "#NestingMode", "#HospitalBag", "#BirthPlan", "#AlmostReady", "#KickCounts", "#PregnancyJourney", "#MamaToBe", "#MommiesMatter"]),
        ("Set 5: Emotional & Community", ["#YouGotThisMama", "#MotherhoodRising", "#PregnancySupport", "#MamaCommunity", "#PregnancyMentalHealth", "#PostpartumPrep", "#NewMomAdvice", "#GentleMotherhood", "#MindfulPregnancy", "#MommiesMatter"]),
    ]

    for set_name, tags in hashtag_sets:
        p = doc.add_paragraph()
        run = p.add_run(set_name)
        run.font.name = "Georgia"
        run.font.size = Pt(14)
        run.font.color.rgb = MOSS_DEEP
        run.font.bold = True
        p = doc.add_paragraph()
        run = p.add_run("  ".join(tags))
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = INK
        doc.add_paragraph()

    doc.add_page_break()

    # Weekly posting rhythm
    add_heading_styled(doc, "Weekly Posting Rhythm", level=2)
    add_para(doc, "A sustainable cadence for solo founders. 4–5 posts per week. Adjust to fit your energy.",
             size=11, italic=True)

    rhythm = [
        ("Monday", "Affirmation", "Weekly affirmation post. Easy to batch-create 4 weeks ahead. Sets a nurturing tone for the week.", "Save + repeat daily."),
        ("Tuesday", "Educational", "Week-specific tip or 'what to expect' post. Use the app's content as source material.", "Save for your week."),
        ("Wednesday", "Engagement", "Ask the audience a question (cravings, week, name ideas). Boosts algorithm reach.", "Comment below."),
        ("Thursday", "Product / BTS", "Feature spotlight, ebook excerpt, or behind-the-scenes story. Sells without selling.", "Link in bio."),
        ("Friday", "Emotional", "Heart-centered post. A reminder, a love letter, an affirmation. Builds brand love.", "Share with a mama."),
        ("Saturday", "Skip", "Rest day. Or repost a UGC mention if any.", ""),
        ("Sunday", "Skip", "Rest day. Plan next week's content.", ""),
    ]

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, header in enumerate(["Day", "Content Type", "What to Post", "CTA"]):
        hdr[i].text = header
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = CREAM
        shade_cell(hdr[i], HEX_MOSS_DEEP)

    for day, ctype, what, cta in rhythm:
        row = table.add_row().cells
        row[0].text = day
        row[1].text = ctype
        row[2].text = what
        row[3].text = cta
        for cell in row:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
                    run.font.color.rgb = INK
        shade_cell(row[0], HEX_CREAM)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(0.8)
        row.cells[1].width = Inches(1.2)
        row.cells[2].width = Inches(4.0)
        row.cells[3].width = Inches(1.5)

    doc.save(os.path.join(OUTPUT_DIR, "social-content-kit.docx"))
    print(f"Generated: {OUTPUT_DIR}/social-content-kit.docx")

# ──────────────────────────────────────────────────────────────────────────
# 2. LAUNCH ANNOUNCEMENT
# ──────────────────────────────────────────────────────────────────────────

def generate_launch_announcement():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FOR IMMEDIATE RELEASE")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = ROSE_GOLD
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Tender Trimesters")
    run.font.name = "Georgia"
    run.font.size = Pt(36)
    run.font.color.rgb = MOSS_DEEP
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("A pregnancy app that feels like a best friend — now live")
    run.font.name = "Georgia"
    run.font.size = Pt(18)
    run.font.color.rgb = MOSS
    run.font.italic = True

    add_horizontal_line(doc)

    # Dateline + lead
    add_para(doc, "AUSTIN, TX — August 15, 2026", size=11, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=MOSS_DEEP)
    add_para(doc, "Mommies Matter, the maternal wellness brand founded by author Helena-Ann Baker, today launched Tender Trimesters — a free pregnancy app designed to feel less like a medical tracker and more like a best friend who happens to know a lot about pregnancy. The app launches alongside Baker's debut book, Mommies Matter: A Gentle Guide for New Mothers, now available on Amazon Kindle.",
             size=12)

    add_para(doc, "Tender Trimesters walks mothers through all 40 weeks of pregnancy with weekly milestone content, daily affirmations, a private journal, mood tracking, and Tempie — a 24/7 AI companion trained to listen without judgment, offer practical tips, and remind mamas they are not alone. The app is free forever; a one-time $9.99 Premium tier unlocks unlimited Tempie chats, bump photo gallery, partner access, audio meditations, and the full Mommies Matter digital bundle.",
             size=12)

    # Helena-Ann quote
    add_para(doc, "", size=8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('"When I was pregnant with my son Tatum, I needed someone to tell me what to actually do. Not the textbook version. The real version. The one that says: eat the saltines, take the nap, call your OB when you\'re worried, and trust that you are not alone in this. So I built it."')
    run.font.name = "Georgia"
    run.font.size = Pt(14)
    run.font.color.rgb = MOSS_DEEP
    run.font.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— Helena-Ann Baker, founder of Mommies Matter and author of Mommies Matter: A Gentle Guide for New Mothers")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = MOSS

    add_horizontal_line(doc)

    # Product details
    add_heading_styled(doc, "About Tender Trimesters", level=2)
    add_para(doc, "Tender Trimesters is a free pregnancy companion app built around four pillars: weekly education, daily emotional support, private reflection, and 24/7 AI companionship. Every feature is designed with one question: does this make a mama feel more held?",
             size=11)

    add_para(doc, "Free tier (forever):", size=11, bold=True, color=MOSS_DEEP)
    features_free = [
        "40-week milestone calendar with baby size comparisons, body changes, emotional guidance, best-friend tips, and self-care checklists for each week",
        "Daily affirmations — one per week of pregnancy, written by hand by Helena-Ann",
        "Mood check-ins with simple trend visualization",
        "Private journal with mood, photo, craving, and baby name fields",
        "Tempie AI chat — 5 messages per day, 24/7",
        "Appointment reminders for OB visits, ultrasounds, and tests",
    ]
    for f in features_free:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = INK

    add_para(doc, "Premium tier ($9.99 one-time or $4.99/month):", size=11, bold=True, color=MOSS_DEEP)
    features_premium = [
        "Unlimited Tempie AI chat — 24/7, no daily cap",
        "Bump photo gallery organized by week",
        "Partner access — a read-only link so partners can follow the journey",
        "Audio affirmations and guided meditations",
        "Kick counter and contraction timer",
        "Letters to Baby journal templates",
        "The full Mommies Matter digital bundle: ebook, 40 printable affirmation cards, First-Trimester Survival Kit, and letters to baby templates",
    ]
    for f in features_premium:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = INK

    # Where to find
    add_heading_styled(doc, "Where to find it", level=2)
    add_para(doc, "Tender Trimesters is available at tendertrimesters.com (link). The Mommies Matter book is available on Amazon Kindle at amazon.com/MommiesMatter (link).",
             size=11)
    add_para(doc, "Press contact: Helena-Ann Baker, hello@mommiesmatter.com",
             size=11)

    add_horizontal_line(doc)

    # About Mommies Matter
    add_heading_styled(doc, "About Mommies Matter", level=2)
    add_para(doc, "Mommies Matter is a maternal wellness brand founded by Helena-Ann Baker in 2025. The company's mission is to make every mother feel seen, supported, and held — through pregnancy, postpartum, and the early years of motherhood. The brand encompasses the Mommies Matter book (Amazon Kindle), the Tender Trimesters app, and a growing library of digital resources for new and expecting mothers. Helena-Ann lives in Austin, TX with her son Tatum and husband.",
             size=11)

    add_horizontal_line(doc)

    # Email nurture sequence
    doc.add_page_break()
    add_heading_styled(doc, "Waitlist Email Nurture Sequence", level=2)
    add_para(doc, "Five emails to send to waitlist subscribers leading up to and after launch. Send via your email platform (Mailchimp, ConvertKit, etc.) with the timing below.",
             size=11, italic=True)

    emails = [
        ("Email 1 — Welcome (send immediately on signup)",
         "Subject: Welcome to Tender Trimesters, mama 💛",
         "Hi {first_name},\n\nWelcome to the Tender Trimesters waitlist. You're officially on the list — and we're so glad you're here.\n\nOver the next 10 days, you'll get a few emails from me. Real ones. Not auto-responder fluff. I'll share my story, give you a peek inside the app, and tell you when we go live.\n\nFor now, here's your first affirmation:\n\n\"My body knows exactly what to do.\"\n\nPin it somewhere. Say it out loud. Repeat as needed.\n\nWith love,\nHelena-Ann"),
        ("Email 2 — Helena-Ann's story (Day 2)",
         "Subject: Why I built this — my story",
         "Hi {first_name},\n\nWhen I was pregnant with my son Tatum, I needed someone to tell me what to actually do. Not the textbook version. The real version.\n\nThe one that says: eat the saltines, take the nap, call your OB when you're worried, and trust that you are not alone in this.\n\nI lost my Granny in 2019. She came to me in a dream and told me I was pregnant before I even knew. She had asked me for years: 'Helena-Ann, when you gonna write your book?' She saw the writer in me before I did.\n\nI believe she met Tatum before I did, and sent him to me because she knew I would need his love after she left.\n\nSo I wrote the book. And then I built the app. For you. For the mama I was.\n\nTomorrow I'll show you inside the app — including Tempie, the AI companion I wish I'd had at 3am.\n\nWith love,\nHelena-Ann"),
        ("Email 3 — Tempie spotlight (Day 4)",
         "Subject: Meet Tempie — your 24/7 AI companion",
         "Hi {first_name},\n\nToday I want to show you inside the app — specifically, the feature I'm most proud of: Tempie.\n\nTempie is an AI companion inside Tender Trimesters. She's awake when your OB is asleep. She's trained to:\n\n• Listen without judgment\n• Offer practical, evidence-informed tips\n• Share affirmations and gentle reminders\n• Help you process emotions at 3am\n• Always defer to your OB for medical decisions\n\nShe's not a doctor. She's a friend who happens to know a lot about pregnancy. And she's there for you, 24/7.\n\nFree tier: 5 messages per day. Premium: unlimited.\n\n3 more days until launch. I can't wait to share this with you.\n\nWith love,\nHelena-Ann"),
        ("Email 4 — Testimonial + book excerpt (Day 7)",
         "Subject: \"It feels like a friend is checking in with me\"",
         "Hi {first_name},\n\nThree more days until we go live. Today, I want to share two things.\n\nFirst, a note from a beta tester:\n\n\"Tempie answered my 3am panic about whether my baby's kicks were normal. She didn't replace my OB, but she helped me breathe until morning. I cried.\"\n\n— Mama of one, week 28\n\nSecond, an excerpt from Mommies Matter (live now on Amazon Kindle):\n\n\"Birth is unpredictable. But that doesn't mean you're powerless. I was 18 hours in. My body was tired. At 7 centimeters, I cried to my husband: 'I can't do it, baby. I need the epidural.' Suddenly the nurses were in motion — like it was a national emergency...\"\n\nRead the full chapter at amazon.com/MommiesMatter.\n\nAlmost there, mama.\n\nWith love,\nHelena-Ann"),
        ("Email 5 — Launch day (Day 10 / Launch Day)",
         "Subject: We're live! 💛",
         "Hi {first_name},\n\nToday's the day. Tender Trimesters is live.\n\nGo to tendertrimesters.com and create your free account. Takes 60 seconds. You'll be inside the app immediately.\n\nWhat's inside:\n\n• 40-week milestone calendar\n• Daily affirmations\n• Mood check-ins\n• Private journal\n• Tempie AI chat (5 msgs/day free)\n• Appointment reminders\n\nWant more? Premium is $9.99 one-time and includes unlimited Tempie, bump photos, partner access, audio meditations, kick counter, and the full Mommies Matter digital bundle.\n\nYou've got this, mama. And we've got you.\n\nWith love,\nHelena-Ann\n\nP.S. If anything feels off or you have feedback, just reply to this email. I read every one."),
    ]

    for subject_line, subject, body in emails:
        p = doc.add_paragraph()
        run = p.add_run(subject_line)
        run.font.name = "Georgia"
        run.font.size = Pt(14)
        run.font.color.rgb = MOSS_DEEP
        run.font.bold = True

        p = doc.add_paragraph()
        run = p.add_run(subject)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = TERRACOTTA
        run.font.bold = True

        p = doc.add_paragraph()
        run = p.add_run(body)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = INK

        add_horizontal_line(doc)

    doc.save(os.path.join(OUTPUT_DIR, "launch-announcement.docx"))
    print(f"Generated: {OUTPUT_DIR}/launch-announcement.docx")

# ──────────────────────────────────────────────────────────────────────────
# 3. INFLUENCER OUTREACH
# ──────────────────────────────────────────────────────────────────────────

def generate_influencer_outreach():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TENDER TRIMESTERS")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = ROSE_GOLD
    run.font.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Influencer & Partner Outreach Templates")
    run.font.name = "Georgia"
    run.font.size = Pt(32)
    run.font.color.rgb = MOSS_DEEP
    run.font.bold = True
    add_para(doc, "Five templates for different partner types  ·  One-page partnership pitch",
             size=11, color=MOSS, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_horizontal_line(doc)
    add_para(doc, "by Mommies Matter  ·  Made with love",
             size=9, color=MOSS, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # How to use
    add_heading_styled(doc, "How to use these templates", level=2)
    add_para(doc, "Each template below is a starting point — not a copy-paste script. Personalize the first line with something specific to the person you're reaching out to (a recent post of theirs, a shared connection, a specific piece of their content you appreciated). Generic outreach gets ignored. Specific outreach gets replies.",
             size=11)
    add_para(doc, "Sending cadence: 5–10 personalized outreach emails per week is more effective than 50 generic ones. Track in a spreadsheet. Follow up once after 7 days. Then let it go.",
             size=11)
    add_para(doc, "Affiliate program: 20% commission on Premium ($9.99 one-time). Ambassador tier (30%) for partners with >10k engaged followers. Custom deals for strategic partners.",
             size=11)

    doc.add_page_break()

    # Templates
    templates = [
        ("Template 1 — Doulas",
         "Subject: A gift for your mamas (free app + 20% partner program)",
         "Hi {first_name},\n\nI've been following your work as a doula — your recent post about {specific post/topic} really resonated. The way you hold space for mamas is exactly what this season needs.\n\nI'm Helena-Ann, founder of Tender Trimesters — a free pregnancy app that feels less like a medical tracker and more like a best friend who happens to know a lot about pregnancy. We just launched alongside my book, Mommies Matter, on Amazon Kindle.\n\nI'd love to offer you and your clients free Premium access (no strings) so you can see if it's something you'd want to recommend. The app includes a 24/7 AI companion named Tempie, weekly milestones for all 40 weeks, mood tracking, journaling, and partner access.\n\nIf you decide to recommend us, we have a 20% affiliate program — but only if it feels aligned. No pressure either way.\n\nWould you like me to send you Premium access?\n\nWith warmth,\nHelena-Ann\n\nP.S. The app is at tendertrimesters.com. The book is on Amazon Kindle."),

        ("Template 2 — OB-GYNs / Midwives",
         "Subject: A free resource for your expecting patients",
         "Dear Dr. {last_name},\n\nI hope this note finds you well. I'm Helena-Ann Baker, founder of Mommies Matter and author of the book by the same name (now on Amazon Kindle).\n\nI'm reaching out because I've built a free pregnancy app called Tender Trimesters that I believe could be a useful supplemental resource for your expecting patients. The app provides:\n\n• Week-by-week pregnancy education aligned with standard milestones\n• A 24/7 AI companion (Tempie) for non-medical emotional support — she's explicitly designed to defer to OBs for medical questions and always tells mamas to call their provider when concerned\n• Mood tracking and private journaling\n• Appointment reminders\n\nThe app is free forever. There's a Premium tier ($9.99) for additional features, but the core educational content is free to all.\n\nI'd be glad to send you a Premium account so you can review it yourself. If you find it valuable, we'd be honored to be included in your patient resources.\n\nThank you for the work you do for mamas.\n\nWarmly,\nHelena-Ann Baker\nFounder, Mommies Matter"),

        ("Template 3 — Pregnancy Influencers",
         "Subject: Collab? Free Premium + 20% affiliate (your mamas will love this)",
         "Hi {first_name},\n\nYour content around {specific topic} has been such a gift to follow — your voice in this space is exactly what mamas need.\n\nI'm Helena-Ann, founder of Tender Trimesters (just launched) and author of Mommies Matter on Amazon Kindle. We built a pregnancy app that's free forever — with a 24/7 AI companion named Tempie, 40 weeks of milestone content, daily affirmations, mood tracking, journaling, and partner access.\n\nI'd love to send you free Premium access to play with. If it resonates and you want to share with your audience, we have:\n\n• 20% affiliate commission on Premium ($9.99 one-time, so $2/sale)\n• Custom content assets if you want to do a Tempie demo or app walkthrough\n• Cross-promotion on our channels (we'll feature your content too)\n\nNo pressure if it's not a fit. Either way, sending love to your growing community.\n\nWith warmth,\nHelena-Ann\n\ntendertrimesters.com  ·  @tendertrimesters"),

        ("Template 4 — Mom Bloggers",
         "Subject: Guest post + affiliate? (free Premium inside)",
         "Hi {first_name},\n\nI've been reading {blog name} for a while — your post on {specific post} was particularly helpful. Thank you for the work you put into this space.\n\nI'm Helena-Ann, founder of Tender Trimesters (a free pregnancy app, just launched) and author of Mommies Matter on Amazon Kindle. I'd love to explore two ways we might collaborate:\n\n1. Guest post: I'd be honored to write a piece for your blog. Topics I love: pregnancy mental health, body image in pregnancy, the 'fourth trimester,' partner communication, postpartum identity. Whatever serves your readers.\n\n2. Affiliate partnership: 20% commission on our Premium tier ($9.99 one-time). I can provide custom graphics, sample social posts, and a free Premium account for your review.\n\nIf either sounds interesting, I'd love to chat. No pressure either way — keep doing what you're doing.\n\nWarmly,\nHelena-Ann Baker\n\ntendertrimesters.com"),

        ("Template 5 — Local Birth Educators",
         "Subject: Workshop bundle for your students? (free Premium + custom co-branded materials)",
         "Hi {first_name},\n\nI loved finding your birth education classes — the way you blend evidence-based info with emotional support is rare and special.\n\nI'm Helena-Ann, founder of Tender Trimesters (a free pregnancy app, just launched) and author of Mommies Matter on Amazon Kindle. I'm reaching out because I think there's a beautiful partnership possible here.\n\nWhat I'd love to explore:\n\n• Free Premium access for all your students (a $9.99 value each, on us)\n• Co-branded workshop materials — we can design a custom affirmation card or checklist with your branding alongside ours\n• Cross-promotion on socials\n• Affiliate: 20% commission on any Premium upgrades from your students\n\nThe app's content complements birth ed perfectly — weekly milestones, mood tracking, journaling, and Tempie (our 24/7 AI companion) for the questions that come up between classes.\n\nWould you be open to a 20-minute call to explore?\n\nWith warmth,\nHelena-Ann Baker\n\ntendertrimesters.com  ·  hello@mommiesmatter.com"),
    ]

    for title, subject, body in templates:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.name = "Georgia"
        run.font.size = Pt(18)
        run.font.color.rgb = MOSS_DEEP
        run.font.bold = True

        p = doc.add_paragraph()
        run = p.add_run(subject)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = TERRACOTTA
        run.font.bold = True
        run.font.italic = True

        p = doc.add_paragraph()
        run = p.add_run(body)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = INK

        add_horizontal_line(doc)

    # Partnership pitch
    doc.add_page_break()
    add_heading_styled(doc, "Mommies Matter — Partnership Pitch (One-Pager)", level=1)
    add_para(doc, "Use this as a one-page overview to send to potential partners after initial contact.",
             size=11, italic=True, color=MOSS)

    add_heading_styled(doc, "Brand Snapshot", level=3)
    add_para(doc, "Mommies Matter is a maternal wellness brand founded in 2025 by author Helena-Ann Baker. We make mamas feel seen, supported, and held — through pregnancy, postpartum, and early motherhood.",
             size=11)

    add_heading_styled(doc, "Our Products", level=3)
    products = [
        ("Mommies Matter book", "Helena-Ann's debut guide for new mothers. 17 chapters covering epidural decisions, premature/late babies, labor realities, her preeclampsia + Tatum birth story, feeding, latching, NICU, sleep, postpartum healing, and more. Live on Amazon Kindle."),
        ("Tender Trimesters app", "Free pregnancy app. 40-week milestone calendar, daily affirmations, mood tracking, private journal, Tempie AI companion (24/7), appointment reminders, partner access. Premium $9.99 one-time or $4.99/mo for unlimited Tempie + bump photos + audio meditations + bundle."),
        ("Digital bundle", "Mommies Matter ebook + 40 affirmation cards (printable) + First-Trimester Survival Kit + Letters to Baby templates. Included with Premium, or available standalone."),
    ]
    for name, desc in products:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{name}: ")
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = MOSS_DEEP
        run = p.add_run(desc)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = INK

    add_heading_styled(doc, "Audience", level=3)
    add_para(doc, "Target: expecting mothers, especially first-time moms, ages 25–40. Predominantly US-based with growing international interest. Audience is values-driven, prioritizes emotional wellness over clinical perfection, and responds to warm/personal content over corporate marketing.",
             size=11)
    add_para(doc, "[Insert current audience stats: email list size, social followers, app downloads, monthly website traffic. Update quarterly.]",
             size=11, italic=True, color=MOSS)

    add_heading_styled(doc, "Partnership Tiers", level=3)
    tiers = [
        ("Affiliate (20% commission)", "Open to all. 20% commission on Premium sales ($9.99 one-time = $2/sale, $4.99/mo = $1/month recurring). 60-day cookie. Self-serve through our partner portal."),
        ("Ambassador (30% commission)", "For partners with >10k engaged followers or established authority in the maternal wellness space. 30% commission, custom content assets, cross-promotion on our channels, early access to new features."),
        ("Strategic Partner (custom)", "For doulas, OBs, birth educators, and established brands. Custom deals including co-branded materials, free Premium for clients/students, dedicated landing pages, revenue share beyond standard tiers."),
    ]
    for name, desc in tiers:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{name}: ")
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = MOSS_DEEP
        run = p.add_run(desc)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = INK

    add_heading_styled(doc, "Why partner with us", level=3)
    reasons = [
        "Brand your audience will trust: warm, nurturing, never salesy. We lead with value, not pitches.",
        "Generous free tier: mamas can use the app forever without paying — your audience gets real value whether they upgrade or not.",
        "Real founder story: Helena-Ann's journey with Tatum and her Granny resonates deeply with mamas. Authentic content travels.",
        "Generous commissions and flexible partnership structures.",
        "Responsive team: we reply to partner emails within 24 hours. Always.",
    ]
    for r in reasons:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(r)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = INK

    add_heading_styled(doc, "Next steps", level=3)
    add_para(doc, "Email hello@mommiesmatter.com to start the conversation. Tell us about your audience and what you're hoping to create. We'll reply within 24 hours.",
             size=11)

    doc.save(os.path.join(OUTPUT_DIR, "influencer-outreach.docx"))
    print(f"Generated: {OUTPUT_DIR}/influencer-outreach.docx")

# ──────────────────────────────────────────────────────────────────────────
# 4. CONTENT CALENDAR (XLSX)
# ──────────────────────────────────────────────────────────────────────────

def generate_content_calendar():
    wb = Workbook()

    # ─── Sheet 1: Overview ───
    ws1 = wb.active
    ws1.title = "30-60-90 Overview"

    # Styles
    title_font = Font(name="Calibri", size=22, bold=True, color="3A4233")
    sub_font = Font(name="Calibri", size=11, italic=True, color="6B7A5A")
    period_font = Font(name="Calibri", size=14, bold=True, color="FFFFFF")
    theme_font = Font(name="Calibri", size=12, bold=True, color="3A4233")
    body_font = Font(name="Calibri", size=11, color="2D2A24")

    # Title
    ws1["A1"] = "Tender Trimesters — 90-Day Content Calendar"
    ws1["A1"].font = title_font
    ws1.merge_cells("A1:E1")

    ws1["A2"] = "Launch date: August 15, 2026  ·  Made with love, mama"
    ws1["A2"].font = sub_font
    ws1.merge_cells("A2:E2")

    # Period headers
    headers = [
        ("A4", "Period", "B4", "Dates", "C4", "Theme", "D4", "Goal", "E4", "Success Metric"),
    ]
    for cell_a, label_a, cell_b, label_b, cell_c, label_c, cell_d, label_d, cell_e, label_e in [headers[0]]:
        ws1[cell_a] = label_a
        ws1[cell_b] = label_b
        ws1[cell_c] = label_c
        ws1[cell_d] = label_d
        ws1[cell_e] = label_e
        for c in [cell_a, cell_b, cell_c, cell_d, cell_e]:
            ws1[c].font = period_font
            ws1[c].fill = PatternFill("solid", fgColor=HEX_MOSS_DEEP)
            ws1[c].alignment = Alignment(horizontal="center", vertical="center")

    periods = [
        ("Days 1–30", "Aug 15 – Sep 13", "LAUNCH + AWARENESS",
         "Get the word out. Drive app downloads. Build social following. Establish the brand voice.",
         "1,000 app signups, 2,000 social followers, 50 waitlist→user conversions"),
        ("Days 31–60", "Sep 14 – Oct 13", "EDUCATION + ENGAGEMENT",
         "Educational content. Build authority. Encourage user-generated content. Promote book + bundle.",
         "2,500 cumulative app signups, 5,000 social followers, 100 Premium conversions, 50 book sales"),
        ("Days 61–90", "Oct 14 – Nov 12", "COMMUNITY + CONVERSION",
         "Lean into community. Showcase testimonials. Push Premium + bundle for holiday prep. Partner outreach ramp-up.",
         "5,000 cumulative app signups, 10,000 social followers, 250 Premium conversions, 5 partner deals signed"),
    ]

    for i, (period, dates, theme, goal, metric) in enumerate(periods):
        row = i + 5
        ws1.cell(row=row, column=1, value=period).font = theme_font
        ws1.cell(row=row, column=2, value=dates).font = body_font
        ws1.cell(row=row, column=3, value=theme).font = theme_font
        ws1.cell(row=row, column=4, value=goal).font = body_font
        ws1.cell(row=row, column=5, value=metric).font = body_font

        # Color-code periods
        bg = [HEX_BLUSH, HEX_SAGE, HEX_BUTTER][i]
        for col in range(1, 6):
            cell = ws1.cell(row=row, column=col)
            cell.fill = PatternFill("solid", fgColor=bg)
            cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)

    # Column widths
    ws1.column_dimensions["A"].width = 14
    ws1.column_dimensions["B"].width = 18
    ws1.column_dimensions["C"].width = 28
    ws1.column_dimensions["D"].width = 45
    ws1.column_dimensions["E"].width = 45
    ws1.row_dimensions[4].height = 30
    for r in range(5, 8):
        ws1.row_dimensions[r].height = 75

    # ─── Sheet 2: Daily Plan ───
    ws2 = wb.create_sheet("Daily Content Plan")

    # Title
    ws2["A1"] = "Daily Content Plan — 90 Days"
    ws2["A1"].font = title_font
    ws2.merge_cells("A1:H1")

    ws2["A2"] = "Starts August 15, 2026  ·  Adjust content types to fit your energy  ·  Rotate hashtags from the Social Content Kit"
    ws2["A2"].font = sub_font
    ws2.merge_cells("A2:H2")

    # Headers
    headers = ["Date", "Day", "Platform", "Content Type", "Topic / Headline", "CTA", "Asset Needed", "Status"]
    for i, h in enumerate(headers, 1):
        cell = ws2.cell(row=4, column=i, value=h)
        cell.font = period_font
        cell.fill = PatternFill("solid", fgColor=HEX_MOSS_DEEP)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws2.row_dimensions[4].height = 30

    # Generate 90 days of content
    import datetime
    start_date = datetime.date(2026, 8, 15)

    # Content rotation pattern (5 days/week of posts, 2 rest days)
    # Each week: Mon=Affirmation, Tue=Educational, Wed=Engagement, Thu=Product/BTS, Fri=Emotional, Sat=Rest, Sun=Rest
    week_pattern = [
        ("Instagram", "Affirmation", "Weekly affirmation card", "Save + repeat daily", "Affirmation graphic"),
        ("TikTok", "Educational", "Week-specific tip (use app content)", "Save for your week", "Talking head video"),
        ("Instagram + TikTok", "Engagement", "Ask the audience a question", "Comment below", "Text-on-screen graphic"),
        ("Instagram", "Product / BTS", "Feature spotlight or behind-the-scenes", "Link in bio", "App screenshot or founder photo"),
        ("Instagram + TikTok", "Emotional", "Heart-centered post / love letter", "Share with a mama", "Founding story or quote graphic"),
        ("—", "REST", "Rest day (or repost UGC)", "—", "—"),
        ("—", "REST", "Plan next week's content", "—", "—"),
    ]

    # Specific topics per week (rotating)
    week_topics = [
        # Days 1-30 (Launch + Awareness)
        ("LAUNCH WEEK: 'We're live!' post", "Free vs Premium comparison", "What's your week? + Tag a mama", "Inside the app: Tempie spotlight", "Helena-Ann's why — founder story"),
        ("Daily affirmation routine", "What to pack in hospital bag", "Cravings check-in", "Inside the app: weekly calendar", "Letter to week 16 you"),
        ("Affirmation: 'My body knows what to do'", "Glucose test — what to expect", "Baby name brainstorm", "Inside the app: mood tracker", "Pregnancy after loss — holding space"),
        ("Affirmation: 'I trust the timing'", "The 5-1-1 rule for contractions", "Where are you in your journey?", "Inside the app: journal feature", "The two-week wait"),
        ("Affirmation: 'Rest is productive'", "First trimester foods to avoid", "Announcement reveal stories", "Inside the app: partner access", "Body image in pregnancy"),
        # Days 31-60 (Education + Engagement)
        ("Monday affirmation", "Anatomy scan — week 20", "Share your bump week", "Mommies Matter book — chapter 1 excerpt", "To the mama spiraling at 3am"),
        ("Monday affirmation", "Kick counts — when to start", "What's your weirdest craving?", "Mommies Matter book — epidural chapter", "The 'fourth trimester' is real"),
        ("Monday affirmation", "First OB visit — what to expect", "Hospital bag must-haves", "Mommies Matter book — feeding chapter", "Asking for help is strategy"),
        ("Monday affirmation", "NT scan explained", "Birth plan preferences", "Inside the app: affirmation cards PDF", "Sleep when baby sleeps (annoying but true)"),
        ("Monday affirmation", "Preeclampsia warning signs", "Postpartum prep", "Mommies Matter book — Helena-Ann's birth story", "You will cry over spilled milk"),
        # Days 61-90 (Community + Conversion)
        ("Monday affirmation", "Group B Strep test", "Mama stories — feature a testimonial", "Premium bundle spotlight", "To the mama who feels behind"),
        ("Monday affirmation", "Signs of labor", "Affirmation that helped you most", "Holiday gift guide for mamas", "The love you didn't know was possible"),
        ("Monday affirmation", "Birth preferences template", "What's one thing you wish you'd known?", "Affirmation cards printable", "Letters to baby — why we write them"),
        ("Monday affirmation", "Postpartum healing timeline", "Tag a mama who needs this", "Mommies Matter bundle as a gift", "Your body is doing extraordinary work"),
        ("Monday affirmation", "Going back to work vs SAHM", "Year-end reflection post", "Premium = the gift that keeps giving", "Made with love, mama — year-end thank you"),
    ]

    for day_offset in range(90):
        date = start_date + datetime.timedelta(days=day_offset)
        weekday_idx = day_offset % 7  # 0=Sat (Aug 15, 2026 is a Saturday)
        # Adjust: Aug 15 2026 is a Saturday. So day 0 = Sat, day 1 = Sun, day 2 = Mon, etc.
        # Map: 0=Sat, 1=Sun, 2=Mon, 3=Tue, 4=Wed, 5=Thu, 6=Fri
        # Our pattern is Mon-Sun: [Mon, Tue, Wed, Thu, Fri, Sat, Sun]
        # So pattern_idx for our weekday:
        # Sat=5, Sun=6, Mon=0, Tue=1, Wed=2, Thu=3, Fri=4
        pattern_map = {0: 5, 1: 6, 2: 0, 3: 1, 4: 2, 5: 3, 6: 4}
        pattern_idx = pattern_map[weekday_idx]
        platform, ctype, default_topic, cta, asset = week_pattern[pattern_idx]

        # Get specific topic for this week
        week_num = day_offset // 7
        if week_num < len(week_topics) and pattern_idx < 5:
            topic = week_topics[week_num][pattern_idx]
        else:
            topic = default_topic

        day_names = ["Saturday", "Sunday", "Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]
        day_name = day_names[weekday_idx]

        row = day_offset + 5
        ws2.cell(row=row, column=1, value=date.strftime("%b %d, %Y")).font = body_font
        ws2.cell(row=row, column=2, value=day_name).font = body_font
        ws2.cell(row=row, column=3, value=platform).font = body_font
        ws2.cell(row=row, column=4, value=ctype).font = body_font
        ws2.cell(row=row, column=5, value=topic).font = body_font
        ws2.cell(row=row, column=6, value=cta).font = body_font
        ws2.cell(row=row, column=7, value=asset).font = body_font
        ws2.cell(row=row, column=8, value="Planned").font = body_font

        # Color code by content type
        type_colors = {
            "Affirmation": HEX_BLUSH,
            "Educational": HEX_SAGE,
            "Engagement": HEX_BUTTER,
            "Product / BTS": "EDE4D1",
            "Emotional": "F4C2C2",
            "REST": "FFFFFF",
        }
        bg = type_colors.get(ctype, "FFFFFF")
        for col in range(1, 9):
            cell = ws2.cell(row=row, column=col)
            cell.fill = PatternFill("solid", fgColor=bg)
            cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)

    # Column widths
    ws2.column_dimensions["A"].width = 14
    ws2.column_dimensions["B"].width = 10
    ws2.column_dimensions["C"].width = 18
    ws2.column_dimensions["D"].width = 16
    ws2.column_dimensions["E"].width = 45
    ws2.column_dimensions["F"].width = 22
    ws2.column_dimensions["G"].width = 25
    ws2.column_dimensions["H"].width = 10

    # Freeze panes
    ws2.freeze_panes = "A5"

    wb.save(os.path.join(OUTPUT_DIR, "content-calendar.xlsx"))
    print(f"Generated: {OUTPUT_DIR}/content-calendar.xlsx")

# ─── Run all ───
if __name__ == "__main__":
    generate_social_content_kit()
    generate_launch_announcement()
    generate_influencer_outreach()
    generate_content_calendar()
    print("\nAll marketing assets generated.")
